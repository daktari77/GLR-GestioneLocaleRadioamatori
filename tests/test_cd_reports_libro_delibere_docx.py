# -*- coding: utf-8 -*-

import os
import tempfile
import unittest
from datetime import date, timedelta
from pathlib import Path

import sys
sys.path.insert(0, str(Path(__file__).parent.parent / 'src'))

from database import exec_query, init_db, set_db_path, fetch_one


class TestCdReportsLibroDelibereDocx(unittest.TestCase):
    def setUp(self):
        self.temp_db = tempfile.NamedTemporaryFile(delete=False, suffix=".db")
        self.temp_db.close()
        self.db_path = self.temp_db.name
        set_db_path(self.db_path)
        init_db()

    def tearDown(self):
        try:
            os.unlink(self.db_path)
        except Exception:
            pass

    def test_export_libro_delibere_docx_writes_content(self):
        import cd_reports
        from docx import Document

        today = date.today()
        past = (today - timedelta(days=10)).isoformat()

        # Minimal template
        with tempfile.NamedTemporaryFile(delete=False, suffix=".docx") as tmp_tpl:
            tpl_path = tmp_tpl.name
        doc = Document()
        doc.add_paragraph("Delibere")
        doc.add_paragraph("2023-2025")
        doc.save(tpl_path)

        try:
            exec_query(
                """
                INSERT INTO cd_riunioni (numero_cd, data, titolo, note, tipo_riunione, created_at)
                VALUES (?, ?, ?, ?, ?, ?)
                """,
                ("1", past, "Riunione", "", None, past + "T00:00:00"),
            )
            cd_id = int(fetch_one("SELECT id FROM cd_riunioni WHERE numero_cd = ?", ("1",))["id"])

            exec_query(
                """
                INSERT INTO cd_delibere (cd_id, numero, oggetto, esito, data_votazione, note, created_at)
                VALUES (?, ?, ?, ?, ?, ?, ?)
                """,
                (cd_id, "01/2026", "Oggetto X", "APPROVATA", past, "Testo delibera", past + "T00:00:00"),
            )

            with tempfile.NamedTemporaryFile(delete=False, suffix=".docx") as tmp_out:
                out_path = tmp_out.name

            try:
                count, warnings = cd_reports.export_libro_delibere_docx(out_path, template_path=tpl_path)
                self.assertEqual(count, 1)
                self.assertIsInstance(warnings, list)

                out_doc = Document(out_path)
                full_text = "\n".join([p.text for p in out_doc.paragraphs if p.text.strip()])
                self.assertIn("delibera n. 01/2026", full_text)
                self.assertIn("Oggetto X", full_text)
                self.assertIn("Testo delibera", full_text)
            finally:
                try:
                    os.unlink(out_path)
                except Exception:
                    pass
        finally:
            try:
                os.unlink(tpl_path)
            except Exception:
                pass


if __name__ == "__main__":
    unittest.main()
