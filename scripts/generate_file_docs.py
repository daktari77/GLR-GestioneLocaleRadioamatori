"""Generate one .docx documentation file per source file.

Strategy (safe + maintainable):
- Prefer existing descriptions from src/V4_MODULE_INDEX.md table.
- Otherwise use the module docstring / header comment.
- Otherwise produce a minimal summary based on discovered symbols (classes/functions).

Outputs mirror the repo-relative paths under docs/file_docs/.
"""

from __future__ import annotations

import argparse
import ast
import csv
import json
import os
import re
from dataclasses import dataclass
from datetime import datetime
from pathlib import Path
from typing import Iterable, Optional

from docx import Document


TEXT_EXTENSIONS = {
    ".py",
    ".md",
    ".ps1",
    ".txt",
    ".spec",
    ".json",
    ".csv",
}

DEFAULT_EXCLUDE_DIRS = {
    ".git",
    ".vs",
    ".vscode",
    ".venv",
    ".venv_old",
    "__pycache__",
    "backup",
    "build",
    "dist",
    "build_pyinstaller",
    "data",
    "data_demo_portable",
    "data_seed_portable",
    "CVS",
    "artifacts",
    "dist_portable",
}


@dataclass(frozen=True)
class FileDoc:
    rel_path: str
    kind: str
    purpose: str
    highlights: list[str]
    mtime_iso: str


def _read_text(path: Path) -> str:
    return path.read_text(encoding="utf-8", errors="replace")


def _is_binary(path: Path) -> bool:
    try:
        with path.open("rb") as f:
            chunk = f.read(8192)
        return b"\x00" in chunk
    except Exception:
        return True


def _parse_module_index(module_index_path: Path) -> dict[str, str]:
    """Parse Purpose from the first markdown table in V4_MODULE_INDEX.md."""
    if not module_index_path.exists():
        return {}
    text = _read_text(module_index_path)

    purpose_map: dict[str, str] = {}
    # Table rows like: | `logger.py` | Logging configuration | `setup_logger()` |
    row_re = re.compile(r"^\|\s*`(?P<mod>[^`]+)`\s*\|\s*(?P<purpose>[^|]+?)\s*\|", re.MULTILINE)
    for m in row_re.finditer(text):
        mod = m.group("mod").strip()
        purpose = m.group("purpose").strip()
        if mod and purpose:
            purpose_map[mod] = purpose
    return purpose_map


def _summarize_python(path: Path, purpose_map: dict[str, str]) -> FileDoc:
    text = _read_text(path)
    rel_name = path.name
    purpose = purpose_map.get(rel_name, "").strip()

    highlights: list[str] = []
    try:
        tree = ast.parse(text)
        doc = ast.get_docstring(tree) or ""
        if not purpose and doc:
            purpose = doc.strip().splitlines()[0].strip()

        for node in tree.body:
            if isinstance(node, ast.ClassDef):
                highlights.append(f"Classe: {node.name}")
            elif isinstance(node, ast.FunctionDef):
                highlights.append(f"Funzione: {node.name}()")
    except Exception:
        if not purpose:
            # fallback: first non-empty comment line
            for line in text.splitlines()[:40]:
                s = line.strip()
                if s.startswith("#"):
                    purpose = s.lstrip("#").strip()
                    break

    if not purpose:
        purpose = "Descrizione non disponibile: documento generato automaticamente."

    mtime_iso = datetime.fromtimestamp(path.stat().st_mtime).isoformat(timespec="seconds")
    return FileDoc(rel_path=str(path), kind="python", purpose=purpose, highlights=highlights[:40], mtime_iso=mtime_iso)


def _summarize_markdown(path: Path) -> FileDoc:
    text = _read_text(path)
    title = ""
    for line in text.splitlines()[:40]:
        s = line.strip()
        if s.startswith("#"):
            title = s.lstrip("#").strip()
            break
    purpose = title or "Documento Markdown del progetto."

    highlights: list[str] = []
    for line in text.splitlines()[:200]:
        s = line.strip()
        if s.startswith("## "):
            highlights.append("Sezione: " + s[3:].strip())

    mtime_iso = datetime.fromtimestamp(path.stat().st_mtime).isoformat(timespec="seconds")
    return FileDoc(rel_path=str(path), kind="markdown", purpose=purpose, highlights=highlights[:40], mtime_iso=mtime_iso)


def _summarize_powershell(path: Path) -> FileDoc:
    text = _read_text(path)

    purpose = "Script PowerShell."
    # comment-based help
    help_match = re.search(r"<#[\s\S]*?#>", text)
    if help_match:
        block = help_match.group(0)
        # take first meaningful line
        for line in block.splitlines():
            s = line.strip("# ").strip()
            if s and not s.startswith("<") and not s.startswith(">"):
                purpose = s
                break

    highlights: list[str] = []
    for m in re.finditer(r"^\s*function\s+([A-Za-z0-9_\-]+)", text, re.MULTILINE):
        highlights.append(f"Funzione: {m.group(1)}")
    for m in re.finditer(r"^\s*param\(\s*$", text, re.MULTILINE):
        highlights.append("Definisce parametri (param(...))")
        break

    mtime_iso = datetime.fromtimestamp(path.stat().st_mtime).isoformat(timespec="seconds")
    return FileDoc(rel_path=str(path), kind="powershell", purpose=purpose, highlights=highlights[:40], mtime_iso=mtime_iso)


def _summarize_json(path: Path) -> FileDoc:
    text = _read_text(path)
    purpose = "File JSON di configurazione/dati."
    highlights: list[str] = []
    try:
        obj = json.loads(text)
        if isinstance(obj, dict):
            keys = list(obj.keys())
            if keys:
                highlights.append("Chiavi principali: " + ", ".join(keys[:25]))
        elif isinstance(obj, list):
            highlights.append(f"Lista con {len(obj)} elementi")
    except Exception:
        pass

    mtime_iso = datetime.fromtimestamp(path.stat().st_mtime).isoformat(timespec="seconds")
    return FileDoc(rel_path=str(path), kind="json", purpose=purpose, highlights=highlights[:40], mtime_iso=mtime_iso)


def _summarize_csv(path: Path) -> FileDoc:
    purpose = "File CSV (dati/import/export)."
    highlights: list[str] = []
    try:
        with path.open("r", encoding="utf-8", errors="replace", newline="") as f:
            reader = csv.reader(f)
            header = next(reader, None)
            if header:
                highlights.append("Colonne: " + ", ".join([h.strip() for h in header if h is not None][:40]))
    except Exception:
        pass

    mtime_iso = datetime.fromtimestamp(path.stat().st_mtime).isoformat(timespec="seconds")
    return FileDoc(rel_path=str(path), kind="csv", purpose=purpose, highlights=highlights[:40], mtime_iso=mtime_iso)


def _summarize_generic(path: Path) -> FileDoc:
    text = _read_text(path)
    purpose = "File di supporto del progetto."
    # first non-empty line
    for line in text.splitlines()[:40]:
        s = line.strip()
        if s:
            purpose = s[:200]
            break

    mtime_iso = datetime.fromtimestamp(path.stat().st_mtime).isoformat(timespec="seconds")
    return FileDoc(rel_path=str(path), kind="text", purpose=purpose, highlights=[], mtime_iso=mtime_iso)


def _summarize_file(path: Path, purpose_map: dict[str, str]) -> Optional[FileDoc]:
    if path.suffix.lower() not in TEXT_EXTENSIONS:
        return None
    if _is_binary(path):
        return None

    suffix = path.suffix.lower()
    if suffix == ".py":
        return _summarize_python(path, purpose_map)
    if suffix == ".md":
        return _summarize_markdown(path)
    if suffix == ".ps1":
        return _summarize_powershell(path)
    if suffix == ".json":
        return _summarize_json(path)
    if suffix == ".csv":
        return _summarize_csv(path)
    return _summarize_generic(path)


def _iter_files(root: Path, exclude_dirs: set[str]) -> Iterable[Path]:
    for dirpath, dirnames, filenames in os.walk(root):
        # prune excluded dirs
        dirnames[:] = [d for d in dirnames if d not in exclude_dirs]
        for fn in filenames:
            yield Path(dirpath) / fn


def _write_docx(doc: FileDoc, repo_root: Path, out_root: Path) -> Path:
    rel = Path(doc.rel_path).resolve().relative_to(repo_root.resolve())
    out_path = out_root / rel
    out_path = out_path.with_name(out_path.name + ".docx")
    out_path.parent.mkdir(parents=True, exist_ok=True)

    d = Document()
    d.add_heading(str(rel), level=1)

    meta = d.add_paragraph()
    meta.add_run("Tipo: ").bold = True
    meta.add_run(doc.kind)
    meta = d.add_paragraph()
    meta.add_run("Ultima modifica: ").bold = True
    meta.add_run(doc.mtime_iso)

    d.add_heading("Funzione", level=2)
    d.add_paragraph(doc.purpose)

    if doc.highlights:
        d.add_heading("Contenuti principali", level=2)
        for h in doc.highlights:
            d.add_paragraph(h, style="List Bullet")

    d.add_paragraph(
        "Nota: documento generato automaticamente; verifica/arricchisci il testo se serve.")

    d.save(out_path)
    return out_path


def _write_index(all_docs: list[tuple[FileDoc, Path]], out_root: Path) -> None:
    d = Document()
    d.add_heading("Indice documentazione file", level=1)
    d.add_paragraph(f"Generato: {datetime.now().isoformat(timespec='seconds')}")

    for doc, out_path in sorted(all_docs, key=lambda t: t[0].rel_path.lower()):
        p = d.add_paragraph(style="List Bullet")
        p.add_run(Path(doc.rel_path).as_posix())
        p.add_run(" — ")
        p.add_run(doc.purpose[:180])
        p.add_run(" (")
        p.add_run(out_path.relative_to(out_root).as_posix())
        p.add_run(")")

    d.save(out_root / "INDEX.docx")


def main() -> int:
    parser = argparse.ArgumentParser(description="Generate .docx documentation per file")
    parser.add_argument("--repo-root", default=str(Path(__file__).resolve().parents[1]))
    parser.add_argument("--out", default="docs/file_docs")
    parser.add_argument("--roots", nargs="+", default=["src", "scripts"])
    args = parser.parse_args()

    repo_root = Path(args.repo_root).resolve()
    out_root = (repo_root / args.out).resolve()
    out_root.mkdir(parents=True, exist_ok=True)

    purpose_map = _parse_module_index(repo_root / "src" / "V4_MODULE_INDEX.md")

    generated: list[tuple[FileDoc, Path]] = []
    for root_rel in args.roots:
        root = (repo_root / root_rel).resolve()
        if not root.exists():
            continue
        for p in _iter_files(root, DEFAULT_EXCLUDE_DIRS):
            fd = _summarize_file(p, purpose_map)
            if fd is None:
                continue
            out_path = _write_docx(fd, repo_root, out_root)
            generated.append((fd, out_path))

    _write_index(generated, out_root)
    print(f"Generated {len(generated)} docx files under: {out_root}")
    return 0


if __name__ == "__main__":
    raise SystemExit(main())
