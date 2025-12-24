# -*- coding: utf-8 -*-
"""Generate a Word document (.docx) for the UI Map.

Usage (from repo root):
  .venv\\Scripts\\python.exe scripts\\generate_ui_map_doc.py

Input:
  docs/revisione/UI_MAP.md

Output:
  docs/revisione/15_UI_Map.docx

Notes:
- This is a best-effort Markdown-to-Word conversion (headings + bullets + paragraphs).
"""

from __future__ import annotations

from datetime import datetime
from pathlib import Path


REPO_ROOT = Path(__file__).resolve().parents[1]
INPUT_MD = REPO_ROOT / "docs" / "revisione" / "UI_MAP.md"
OUTPUT_DOCX = REPO_ROOT / "docs" / "revisione" / "15_UI_Map.docx"


def _safe_doc_save(doc, out_path: Path) -> Path:
    try:
        doc.save(out_path)
        return out_path
    except PermissionError:
        suffix = datetime.now().strftime("%Y%m%d_%H%M%S")
        alt_path = out_path.with_name(f"{out_path.stem}_{suffix}{out_path.suffix}")
        doc.save(alt_path)
        return alt_path


def _parse_md_lines(text: str):
    """Yield tuples (kind, value) where kind is 'h1'|'h2'|'h3'|'bullet'|'para'|'blank'."""
    for raw in text.splitlines():
        line = raw.rstrip("\n")
        stripped = line.strip()
        if not stripped:
            yield ("blank", "")
            continue
        if stripped.startswith("### "):
            yield ("h3", stripped[4:].strip())
            continue
        if stripped.startswith("## "):
            yield ("h2", stripped[3:].strip())
            continue
        if stripped.startswith("# "):
            yield ("h1", stripped[2:].strip())
            continue
        if stripped.startswith("- "):
            yield ("bullet", stripped[2:].strip())
            continue
        yield ("para", stripped)


def generate_ui_map_docx() -> Path:
    try:
        from docx import Document  # type: ignore
    except Exception as exc:
        raise SystemExit(
            "python-docx non disponibile. Installa con: pip install python-docx\n"
            f"Dettagli: {exc}"
        )

    if not INPUT_MD.exists():
        raise SystemExit(f"File UI map non trovato: {INPUT_MD}")

    text = INPUT_MD.read_text(encoding="utf-8", errors="replace")

    doc = Document()

    # Render markdown-ish structure
    pending_blank = False
    for kind, value in _parse_md_lines(text):
        if kind == "blank":
            pending_blank = True
            continue

        if kind == "h1":
            if pending_blank:
                pending_blank = False
            doc.add_heading(value, level=0)
            continue
        if kind == "h2":
            if pending_blank:
                pending_blank = False
            doc.add_heading(value, level=1)
            continue
        if kind == "h3":
            if pending_blank:
                pending_blank = False
            doc.add_heading(value, level=2)
            continue

        if kind == "bullet":
            doc.add_paragraph(value, style="List Bullet")
            pending_blank = False
            continue

        # paragraph
        if pending_blank:
            # create a visual separation without excessive empty lines
            doc.add_paragraph("")
            pending_blank = False
        doc.add_paragraph(value)

    return _safe_doc_save(doc, OUTPUT_DOCX)


def main() -> int:
    out_path = generate_ui_map_docx()
    print(str(out_path))
    return 0


if __name__ == "__main__":
    raise SystemExit(main())
