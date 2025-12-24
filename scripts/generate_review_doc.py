# -*- coding: utf-8 -*-
"""Generate Word documents (docx) for GestioneSoci functional review.

Usage examples (from repo root):
  .venv\\Scripts\\python.exe scripts\\generate_review_doc.py --section 01
    .venv\\Scripts\\python.exe scripts\\generate_review_doc.py --combined

Notes:
- Output is .docx (Word). If you need legacy .doc, open in Word/LibreOffice and Save As .doc.
"""

from __future__ import annotations

import argparse
from dataclasses import dataclass
from datetime import date, datetime
from pathlib import Path
import re
from typing import Iterable, List


@dataclass(frozen=True)
class Section:
    code: str
    title: str
    file_stem: str


REPO_ROOT = Path(__file__).resolve().parents[1]
OUTPUT_DIR = REPO_ROOT / "docs" / "revisione"


def _add_heading(doc, text: str, level: int = 1):
    doc.add_heading(text, level=level)


def _add_paragraph(doc, text: str):
    doc.add_paragraph(text)


def _add_centered_paragraph(doc, text: str, style: str | None = None):
    """Add a centered paragraph, optionally with a Word style."""

    p = doc.add_paragraph(text)
    if style:
        try:
            p.style = style
        except Exception:
            # Style may not exist in some environments; keep default.
            pass
    try:
        from docx.enum.text import WD_ALIGN_PARAGRAPH  # type: ignore

        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    except Exception:
        # If enum import fails, leave default alignment.
        pass
    return p


def _read_version_build() -> tuple[str | None, str | None]:
    """Best-effort extraction of APP_VERSION and BUILD_ID from src/config.py."""

    cfg_path = REPO_ROOT / "src" / "config.py"
    if not cfg_path.exists():
        return None, None
    try:
        text = cfg_path.read_text(encoding="utf-8", errors="replace")
    except Exception:
        return None, None

    m_ver = re.search(r"^\s*APP_VERSION\s*=\s*['\"]([^'\"]+)['\"]\s*$", text, re.MULTILINE)
    m_build = re.search(r"^\s*BUILD_ID\s*=\s*['\"]([^'\"]+)['\"]\s*$", text, re.MULTILINE)
    return (m_ver.group(1) if m_ver else None), (m_build.group(1) if m_build else None)


def _safe_doc_save(doc, out_path: Path) -> Path:
    """Save a docx handling Windows/Office file locks with a fallback filename."""

    try:
        doc.save(out_path)
        return out_path
    except PermissionError:
        suffix = datetime.now().strftime("%Y%m%d_%H%M%S")
        alt_path = out_path.with_name(f"{out_path.stem}_{suffix}{out_path.suffix}")
        doc.save(alt_path)
        return alt_path


def _add_bullets(doc, items: Iterable[str]):
    for item in items:
        doc.add_paragraph(item, style="List Bullet")


def _section_01_content() -> dict:
    today = date.today().isoformat()

    return {
        "titolo": "Panoramica e Avvio",
        "sottotitolo": f"Revisione funzionale – generato il {today}",
        "blocchi": [
            (
                "Scopo",
                [
                    "Questo documento descrive le funzionalità e il flusso di avvio del gestionale GestioneSoci (GLR – Gestione Locale Radioamatori), con focus sugli aspetti utili ad una revisione: dati coinvolti, percorsi, moduli, controlli e punti di attenzione.",
                    "La descrizione è orientata all’uso e alla manutenzione (non è una API reference).",
                ],
            ),
            (
                "Cosa gestisce il gestionale (macro-aree)",
                [
                    "Anagrafica soci (attivi/disattivi), con ricerca e gestione campi principali.",
                    "Documenti per socio (archiviazione su filesystem + tracciamento su DB).",
                    "Documenti di sezione (catalogo categorie, metadati e indice).",
                    "Consiglio Direttivo: riunioni, delibere, verbali (con allegati).",
                    "Email Wizard (creazione bozza mail, esportazione .eml, integrazione con client).",
                    "Backup e manutenzione (backup all’avvio, backup manuale, verifiche DB).",
                    "Domini aggiuntivi v4.2: Magazzino/Prestiti, Ponti/Ripetitori, Assicurazioni.",
                ],
            ),
            (
                "Struttura dati (cartelle principali)",
                [
                    "data/\n  - soci.db (SQLite)\n  - documents/ (documenti per socio, per matricola)\n  - section_docs/ (documenti di sezione, catalogati)\n  - .trash/ (cestino documenti)\n  - logs/ (log applicazione)",
                    "backup/ (cartella backup; include backup automatici e ZIP on-demand)",
                ],
            ),
            (
                "Avvio applicazione (sequenza ad alto livello)",
                [
                    "Entry point: src/main.py (avvio da sviluppo: `cd src; python main.py`).",
                    "Config iniziale: definizione percorsi e costanti in src/config.py.",
                    "Setup logger: src/logger.py (log su file in data/logs/app.log).",
                    "Dependency injection via setter: set_db_path(), set_causali_path(), set_config_paths(), set_docs_base().",
                    "Init DB: init_db() crea/aggiorna schema se mancante.",
                    "Backup all’avvio: backup_on_startup(DB_NAME, BACKUP_DIR).",
                    "Splash/Loading: LoadingWindow mostra avanzamento e log in tempo reale.",
                    "Startup checks: collect_startup_issues() rileva problemi (es. documenti DB non presenti su disco).",
                    "UI principale: App(startup_issues=...) crea root Tk, carica config e costruisce i pannelli.",
                    "Wizard primo avvio: App._check_first_run() mostra la procedura guidata se manca la configurazione sezione.",
                ],
            ),
            (
                "Configurazione e pattern importanti",
                [
                    "Pattern fondamentale: in src/main.py i setter vanno chiamati prima di importare UI/moduli che leggono configurazioni (evita errori ‘Config not set’ e path errati).",
                    "Percorsi: in modalità portable (PyInstaller) BASE_DIR è la root della cartella portable; in modalità sviluppo BASE_DIR è src/.",
                    "Date: formato ISO YYYY-MM-DD (funzioni helper in src/utils.py).",
                    "DB: usare placeholder `?` nei parametri SQL; f-string solo per nomi colonna (come da guideline).",
                ],
            ),
            (
                "Controlli all’avvio (startup_checks)",
                [
                    "Verifica presenza file documenti: confronta i record in tabella documenti con l’esistenza del file su disco.",
                    "In caso di file mancanti: produce una lista di issue (ID documento, nominativo, nome file, percorso registrato).",
                    "Le issue vengono passate alla UI per notifica/gestione (in base a come il pannello le visualizza).",
                ],
            ),
            (
                "Output e artefatti di diagnosi",
                [
                    "Log applicazione: data/logs/app.log.",
                    "Backup automatici: backup/ (politica ‘ultimi N’ gestita dal modulo backup).",
                    "Documenti cestinati: data/.trash/.",
                ],
            ),
        ],
    }


def _section_02_content() -> dict:
    today = date.today().isoformat()

    return {
        "titolo": "Anagrafica Soci",
        "sottotitolo": f"Revisione funzionale – generato il {today}",
        "blocchi": [
            (
                "Scopo",
                [
                    "Questa sezione descrive la gestione dell’anagrafica soci: visualizzazione elenco, filtri/ricerca, inserimento e modifica tramite form, eliminazione (soft delete) e gestione cestino.",
                ],
            ),
            (
                "UI: Tab ‘Soci’ (componenti e comandi)",
                [
                    "Toolbar azioni: Nuovo, Modifica, Elimina, 🗑️ Cestino.",
                    "Toolbar strumenti: 📄 Documentale (documenti del socio), Modifica campi (batch edit), Esporta.",
                    "Ricerca: campo ‘Cerca’ con aggiornamento live + pulsante Reset.",
                    "Filtri: Privacy (tutti / senza privacy / con privacy), Stato (tutti / attivi / inattivi), Dati (tutti / dati mancanti / completi).",
                    "Tabella soci (Treeview): prima colonna ‘⚠’ per warning; righe colorate con tag (active/inactive/no_privacy/missing_data).",
                    "Form socio (MemberForm) sotto la tabella, con pulsanti Salva/Annulla allineati alle Note.",
                    "Anteprima documenti del socio: mini-tabella ‘Documenti del socio’ sotto il form (solo informativa rapida).",
                ],
            ),
            (
                "Campi del form (MemberForm)",
                [
                    "Identificazione: Matricola, Nominativo, Nominativo 2, Familiare, Tipo socio (HAM/RCL/THR).",
                    "Contatti: Email, Telefono.",
                    "Dati anagrafici: Nome, Cognome, Codice Fiscale, Data nascita, Luogo nascita.",
                    "Residenza: Indirizzo, CAP, Città, Provincia.",
                    "Stato: Attivo (checkbox), Voto (checkbox), Privacy (checkbox), Q0/Q1/Q2 (codici quota), Ruoli multipli (tag editor).",
                    "Note: campo testo multi-linea compatto.",
                ],
            ),
            (
                "Validazioni lato UI (prima del salvataggio)",
                [
                    "Nominativo: obbligatorio; eccezione per socio tipo RCL (se vuoto, viene salvato ‘-’).",
                    "Email: controllo minimale (se valorizzata deve contenere ‘@’).",
                    "Matricola: se valorizzata deve essere univoca tra i soci non cestinati; se esiste in cestino viene richiesto di ripristinare o usare matricola diversa.",
                    "Data nascita: accetta DD/MM/YYYY o YYYY-MM-DD (validazione tramite ddmmyyyy_to_iso).",
                    "Tipo socio: se valorizzato deve essere uno tra HAM, RCL, THR.",
                    "Nota: esiste anche una validazione più forte nel modello src/models.py (email/CF/CAP/provincia/date/quote), usata tipicamente da import e test; la UI applica solo controlli essenziali.",
                ],
            ),
            (
                "Salvataggio (insert/update)",
                [
                    "Il pulsante Salva chiama App._save_member(): se current_member_id è valorizzato fa UPDATE, altrimenti INSERT.",
                    "I booleani nel form vengono salvati come stringhe '1'/'0' (checkbox).",
                    "Ruoli: il form espone ‘roles’ (lista) e ‘cd_ruolo’ (ruolo primario). Dopo insert/update, set_member_roles() sincronizza soci_ruoli e soci.cd_ruolo.",
                    "Dopo il salvataggio vengono riapplicati i filtri correnti (non un refresh ‘grezzo’) per preservare lo stato di ricerca/filtri.",
                ],
            ),
            (
                "Eliminazione e cestino",
                [
                    "Elimina socio (toolbar/menu): esegue un soft delete impostando soci.deleted_at = datetime('now').",
                    "Cestino: mostra i soci eliminati; consente Ripristina (deleted_at = NULL), Elimina definitivamente (DELETE FROM soci WHERE id=?), e Svuota cestino (DELETE WHERE deleted_at IS NOT NULL).",
                    "L’eliminazione definitiva è irreversibile e può avere effetti a cascata via FK (es. documenti associati con ON DELETE CASCADE).",
                ],
            ),
            (
                "Persistenza (DB): tabella soci e vincoli rilevanti",
                [
                    "Tabella soci: campi principali (matricola, nominativo, nome/cognome, contatti, stato, privacy, quote, note, deleted_at).",
                    "Indice univoco: matricola è UNIQUE quando non NULL (UX parziale); la UI fa anche un controllo applicativo prima del salvataggio.",
                    "Ruoli multipli: tabella soci_ruoli + colonna soci.cd_ruolo (ruolo primario).",
                ],
            ),
            (
                "Punti di attenzione per la revisione",
                [
                    "Formato date: la UI valida, ma non forza la conversione automatica; verificare policy desiderata (salvare sempre ISO o accettare anche DD/MM/YYYY).",
                    "Validazioni: il modello Member è più rigoroso della UI; allineare se necessario (es. CF, CAP, provincia).",
                    "Prestazioni: _check_missing_data() può fare query aggiuntive per telefono/indirizzo per ogni riga; su DB grandi può incidere sul refresh/elenco.",
                ],
            ),
        ],
    }


def _section_03_content() -> dict:
    today = date.today().isoformat()

    return {
        "titolo": "Documenti Socio (Documentale)",
        "sottotitolo": f"Revisione funzionale – generato il {today}",
        "blocchi": [
            (
                "Scopo",
                [
                    "Questa sezione descrive la gestione dei documenti associati ai soci: caricamento file, metadati (categoria/descrizione), apertura e cancellazione, gestione privacy, e differenze tra viste UI (Documentale per socio vs pannello ‘Documenti soci’).",
                ],
            ),
            (
                "Concetti chiave",
                [
                    "Un documento è: file su disco + record DB in tabella documenti.",
                    "Campi DB principali: socio_id, nome_file, percorso, tipo, categoria, descrizione, data_caricamento.",
                    "Categorie: elenco controllato (catalogo) in documents_catalog.py; fallback alla categoria di default.",
                    "Tipi: almeno ‘documento’ e ‘privacy’ (usati per filtri e badge privacy).",
                ],
            ),
            (
                "Dove si usa (punti UI)",
                [
                    "Tab ‘Soci’ → pulsante ‘📄 Documentale’: apre DocumentsDialog per il socio selezionato.",
                    "Tab ‘Documenti’ → sottotab ‘Documenti soci’: DocumentPanel può mostrare tutti i documenti e filtrare per socio/categoria/testo.",
                    "Anteprima compatta in basso nel tab ‘Soci’: mostra i documenti del socio selezionato (lettura rapida).",
                ],
            ),
            (
                "DocumentsDialog (per singolo socio)",
                [
                    "Filtri: radiobutton per tipo documento (tutti/privacy/documenti/altro).",
                    "Azioni: Carica documento, Carica privacy, Visualizza, Modifica metadati, Elimina, Apri cartella.",
                    "Metadati: richiesti tramite popup (categoria + descrizione).",
                    "Shortcut: Enter (visualizza), Delete (elimina), Ctrl+U (carica documento), Ctrl+P (carica privacy), Esc (chiudi).",
                ],
            ),
            (
                "Caricamento documenti (logica)",
                [
                    "Il caricamento usa documents_manager.upload_document(socio_id, file_path, doc_type, categoria, descrizione).",
                    "Destinazione: data/documents/<token_socio>/<doc_type>/ con filename generato (token esadecimale + estensione originale).",
                    "Persistenza DB: add_documento() salva nome_file (il token), percorso (path assoluto/relativo secondo chiamante), tipo, categoria, descrizione e data_caricamento.",
                    "Dopo upload: viene aggiornato un indice testuale per socio (elenco_documenti.txt) tramite refresh interno.",
                ],
            ),
            (
                "Privacy (documento + flag socio)",
                [
                    "Carica privacy: salva il file come tipo ‘privacy’ e imposta privacy_signed=1 sul socio.",
                    "All’apertura del dialog: badge privacy mostra lo stato da DB; se esiste un documento privacy ma privacy_signed risulta falso, il dialog può riallineare impostando privacy_signed=1.",
                ],
            ),
            (
                "Visualizzazione / apertura file",
                [
                    "Apertura documento: tenta apertura con il viewer di default (os.startfile su Windows).",
                    "Apri cartella: apre la directory dei documenti del socio (basata su token socio).",
                    "Nelle viste aggregate, i file mancanti su disco vengono evidenziati (tag ‘missing’ + info ‘File mancante’).",
                ],
            ),
            (
                "Modifica metadati (categoria/descrizione)",
                [
                    "Modifica: aggiorna categoria e/o descrizione sul record DB (non rinomina file).",
                    "Dopo modifica: rigenera l’indice documenti per il socio.",
                ],
            ),
            (
                "Eliminazione documenti",
                [
                    "Elimina: rimuove il file da disco (os.remove) e poi elimina il record DB (DELETE FROM documenti).",
                    "Nota: a differenza dei soci, qui non c’è soft-delete; la cancellazione è definitiva.",
                ],
            ),
            (
                "Differenze tra implementazioni (nota tecnica)",
                [
                    "Esistono due percorsi storici: documents_manager.py (token per cartella socio) e documents.py/utils.docs_dir_for_matricola (cartella per matricola).",
                    "La UI moderna del ‘Documentale’ usa documents_manager.py; scripts esterni/legacy possono usare documents.py.",
                    "Punto di revisione: standardizzare la convenzione di storage (token vs matricola) per ridurre ambiguità nei percorsi e nei controlli di startup.",
                ],
            ),
        ],
    }


def _section_04_content() -> dict:
    today = date.today().isoformat()

    return {
        "titolo": "Documenti di Sezione",
        "sottotitolo": f"Revisione funzionale – generato il {today}",
        "blocchi": [
            (
                "Scopo",
                [
                    "Questa sezione descrive la gestione dei documenti ‘di sezione’ (non legati a un singolo socio): archiviazione per categorie, metadata, indice testuale e operazioni UI (carica/apri/modifica/elimina).",
                ],
            ),
            (
                "Storage su filesystem e categorie",
                [
                    "Root: data/section_docs/ (configurato via SEC_DOCS).",
                    "Categorie predefinite: Verbali CD, Bilanci, Regolamenti, Modulistica, Documenti ARI, Quote ARI, Altro.",
                    "Per ogni categoria esiste una sottocartella ‘slug’ (es. verbali_cd/, bilanci/, …).",
                    "Ogni documento viene copiato nella categoria scelta con filename interno basato su token (hash_id) + estensione originale.",
                ],
            ),
            (
                "Metadata: metadata.json",
                [
                    "I metadata sono persistiti in data/section_docs/metadata.json.",
                    "Schema: {schema_version, documents:{hash_id:{categoria, original_name, stored_name, description, uploaded_at, relative_path}}}.",
                    "Caricamento robusto: se il file è assente o corrotto, il sistema riparte con metadata vuoti (warning su log).",
                    "Le operazioni principali (add/update/delete) mantengono coerenti file + metadata.",
                ],
            ),
            (
                "Indice categoria: elenco_documenti.txt",
                [
                    "Per ogni categoria può essere (ri)generato elenco_documenti.txt (tab-separated): Nome file originale, Descrizione, Categoria, Percorso relativo.",
                    "Generazione: ensure_section_index_file(categoria) rigenera sempre il file quando invocato.",
                    "Uso pratico: allegare/stampare un indice ‘leggibile’ senza aprire l’app.",
                ],
            ),
            (
                "Operazioni supportate (logica)",
                [
                    "Aggiungi: add_section_document(source_path, categoria, descrizione) copia il file e crea entry metadata.",
                    "Modifica: update_section_document_metadata(path, categoria, descrizione) aggiorna description e sposta il file tra cartelle se cambia categoria.",
                    "Elimina: delete_section_document(path) rimuove il file e ripulisce i metadata associati.",
                    "Utility: rename_section_documents_to_schema() per rinominare documenti non indicizzati secondo schema SEZIONE_<CATEGORIA>_<DATA> (solo per file fuori metadata).",
                ],
            ),
            (
                "UI: SectionDocumentPanel (tab Documenti)",
                [
                    "Mostra lista documenti di sezione con filtro per categoria e ricerca testuale.",
                    "Azioni: Carica (file dialog), Apri file, Apri cartella categoria, Modifica metadata, Elimina.",
                    "Quando si apre una cartella categoria, viene rigenerato l’indice elenco_documenti.txt per allineare contenuti/metadata.",
                ],
            ),
        ],
    }


def _section_05_content() -> dict:
    today = date.today().isoformat()

    return {
        "titolo": "Backup, Verifiche e Ripristino",
        "sottotitolo": f"Revisione funzionale – generato il {today}",
        "blocchi": [
            (
                "Scopo",
                [
                    "Questa sezione descrive il sistema di backup del DB SQLite: backup incrementale all’avvio, backup manuale ‘completo’ in ZIP, verifica integrità e ripristino con safety-backup.",
                ],
            ),
            (
                "Backup incrementale (all’avvio)",
                [
                    "Entry: backup_on_startup(db_name, backup_dir) invocato in fase bootstrap.",
                    "Meccanismo: backup_incremental() calcola hash SHA256 del DB e crea un nuovo .db solo se il DB è cambiato dall’ultimo backup.",
                    "Prima di salvare: PRAGMA integrity_check sul DB (se fallisce, il backup viene bloccato).",
                    "Retention: mantiene gli ultimi N backup (default 20) eliminando i più vecchi.",
                    "Metadata: .backup_meta.json (last_backup_hash, last_backup_time, last_backup_file).",
                ],
            ),
            (
                "Backup manuale (on-demand) in ZIP",
                [
                    "Funzione: backup_on_demand(data_dir, db_path, backup_dir) crea un archivio ZIP contenente ‘data/’ (escludendo il DB se già incluso) + copia consistente del DB.",
                    "Copia DB consistente: usa l’API sqlite3.backup per evitare snapshot incoerenti.",
                    "Manifest: backup_manifest.json include created_at, sorgenti, hash del DB e contenuti archiviati.",
                    "Pulizia: la cartella temporanea viene rimossa dopo la creazione dello ZIP.",
                ],
            ),
            (
                "Ripristino da backup",
                [
                    "Funzione: restore_from_backup(backup_path, target_db_path, create_safety_backup=True).",
                    "Verifica integrità sul file backup prima del restore.",
                    "Safety backup: se esiste un DB target, viene creata una copia *.pre_restore_<timestamp> prima di sovrascrivere.",
                    "Post-restore: esegue integrity_check sul DB ripristinato; se fallisce tenta revert al safety backup.",
                ],
            ),
            (
                "Manutenzione",
                [
                    "verify_db(db_name): wrapper per integrità DB.",
                    "rebuild_indexes(db_name): ricrea gli indici definiti in database.CREATE_INDEXES (best-effort).",
                ],
            ),
        ],
    }


def _section_06_content() -> dict:
    today = date.today().isoformat()

    return {
        "titolo": "Email Wizard",
        "sottotitolo": f"Revisione funzionale – generato il {today}",
        "blocchi": [
            (
                "Scopo",
                [
                    "Il wizard email aiuta a preparare comunicazioni verso i soci (o verso il Consiglio Direttivo) usando template e destinatari estratti dal DB. Supporta mailto, esportazione .eml e apertura bozza in Thunderbird.",
                ],
            ),
            (
                "Destinatari (estrazione da DB)",
                [
                    "Modalità ‘Soci attivi’: tutti i soci attivi con email valorizzata.",
                    "Modalità ‘Consiglio Direttivo’: soci attivi con cd_ruolo valorizzato e diverso da ‘Socio’/‘Ex Socio’, con email.",
                    "I destinatari vengono messi in BCC (privacy) e viene mostrato un contatore (n destinatari).",
                ],
            ),
            (
                "Template e ODG",
                [
                    "Template disponibili: Convocazione CD, Comunicazione generale, Convocazione assemblea, Promemoria quota, Personalizzata.",
                    "Il corpo email può includere placeholder {odg}; l’ODG può essere incollato o caricato da una riunione CD selezionata (carica campo ‘odg’ della riunione).",
                ],
            ),
            (
                "Output: mailto vs .eml vs Thunderbird",
                [
                    "Crea Email: genera un URL mailto con subject/body/bcc e apre il client predefinito; se l’URL supera una soglia, propone copia negli appunti.",
                    "Salva .eml: crea un file EML con Subject/To/Bcc/Date e testo; default directory: SEC_DOCS/email_eml/.",
                    "Invia con Thunderbird: usa -compose (subject/body/bcc) e richiede un percorso exe configurato in Preferenze.",
                ],
            ),
            (
                "Email salvate",
                [
                    "Tab dedicata che elenca i .eml salvati; consente Apri, Elimina, Apri cartella, Avvia Thunderbird.",
                    "Selezionando un .eml, il wizard può estrarre Subject/Body/Bcc dal file e usarli per la bozza in Thunderbird.",
                ],
            ),
        ],
    }


def _section_07_content() -> dict:
    today = date.today().isoformat()

    return {
        "titolo": "Preferenze e Configurazione Sezione",
        "sottotitolo": f"Revisione funzionale – generato il {today}",
        "blocchi": [
            (
                "Scopo",
                [
                    "Questa sezione descrive la configurazione persistente della sezione (dati anagrafici sezione, membri CD, ecc.) e alcune preferenze applicative (stati socio, percorso Thunderbird).",
                ],
            ),
            (
                "Preferenze (PreferencesDialog)",
                [
                    "Tab ‘Sezione’: campi come nome/codice sezione, sedi, contatti, coordinate bancarie, recapiti, lista componenti CD (multi-line).",
                    "Tab ‘Stato socio’: consente definire voci personalizzate per lo stato/ruolo (aggiunte alle voci predefinite).",
                    "Tab ‘Client posta’: percorso Thunderbird Portable (se vuoto usa il default di config.py).",
                ],
            ),
            (
                "Persistenza configurazione",
                [
                    "Lettura/scrittura tramite config_manager.load_config() / save_config().",
                    "Le preferenze salvano: custom_role_options, thunderbird_path e i campi sezione.",
                ],
            ),
            (
                "Primo avvio (First Run Wizard)",
                [
                    "Se manca configurazione minima, all’avvio la UI può proporre la procedura guidata di primo avvio.",
                    "Obiettivo: inizializzare i parametri sezione necessari a stampe/export e ad alcune funzioni operative.",
                ],
            ),
        ],
    }


def _section_08_content() -> dict:
    today = date.today().isoformat()

    return {
        "titolo": "Import/Export (CSV)",
        "sottotitolo": f"Revisione funzionale – generato il {today}",
        "blocchi": [
            (
                "Scopo",
                [
                    "Questa sezione descrive i flussi di importazione ed esportazione dati in formato CSV: soci e magazzino.",
                ],
            ),
            (
                "Launcher unificati",
                [
                    "UnifiedImportWizard: permette di scegliere tra import CSV Soci o Magazzino.",
                    "UnifiedExportWizard: permette di scegliere tra export Soci o Magazzino.",
                ],
            ),
            (
                "Export Soci (ExportDialog)",
                [
                    "Selezione campi esportati tramite checkbox (preset: tutto/nessuno/essenziali).",
                    "Filtri: includi soci inattivi; includi soci eliminati (cestino).",
                    "Output: file CSV scelto dall’utente; query dinamica SELECT <campi> FROM soci ... ORDER BY nominativo.",
                    "Normalizzazione output: campi boolean-like (attivo, voto) esportati come ‘Si/No’.",
                ],
            ),
            (
                "Import Soci (ImportWizard)",
                [
                    "Wizard a step: selezione file, configurazione mapping e campi, anteprima, importazione.",
                    "Supporta auto-rilevazione delimitatore e mapping assistito.",
                    "Consente scegliere quali campi aggiornare (✓ = campo aggiornato durante l’import).",
                    "Gestisce casi di conflitto (aggiorna solo vuoti / sovrascrivi / interrompi) con richiesta all’utente.",
                ],
            ),
            (
                "Import/Export Magazzino",
                [
                    "Import: tramite dialog dedicato (supporto CSV/Excel), con logica di upsert su numero inventario.",
                    "Export: dialog dedicato con filtri (disponibili vs prestati) e campi configurabili.",
                ],
            ),
        ],
    }


def _section_09_content() -> dict:
    today = date.today().isoformat()

    return {
        "titolo": "Consiglio Direttivo (Riunioni, Delibere, Verbali)",
        "sottotitolo": f"Revisione funzionale – generato il {today}",
        "blocchi": [
            (
                "Scopo",
                [
                    "Questa sezione descrive la gestione del Consiglio Direttivo: registro riunioni, delibere e verbali con possibilità di associare documenti (doc/docx/pdf) tramite path.",
                ],
            ),
            (
                "Riunioni CD (cd_riunioni)",
                [
                    "Operazioni: elenco (get_all_meetings), dettaglio (get_meeting_by_id), aggiungi (add_meeting), modifica (update_meeting), elimina (delete_meeting).",
                    "Dati principali: numero_cd, data (ISO), titolo, odg/note, verbale_path.",
                    "Validazione allegato: validate_verbale_file accetta .doc/.docx/.pdf.",
                ],
            ),
            (
                "Delibere (cd_delibere)",
                [
                    "Operazioni: get_all_delibere (anche per meeting), get_delibera_by_id, add/update/delete.",
                    "Campi: numero (es. 1/2025), oggetto, esito (APPROVATA/RESPINTA/RINVIATA), data_votazione, favorevoli/contrari/astenuti, allegato_path, note.",
                ],
            ),
            (
                "Verbali (cd_verbali)",
                [
                    "Operazioni: get_all_verbali (anche per meeting), get_verbale_by_id, add/update/delete.",
                    "Campi: data_redazione, segretario, presidente, odg, documento_path, note.",
                    "Validazione documento: validate_documento accetta .doc/.docx/.pdf.",
                ],
            ),
            (
                "UI",
                [
                    "La UI principale offre dialog dedicati per lista e modifica (MeetingsListDialog/MeetingDialog, DeliberaDialog, VerbaleDialog).",
                    "La sezione Email può caricare l’ODG da una riunione CD per generare convocazioni.",
                ],
            ),
        ],
    }


def _section_10_content() -> dict:
    today = date.today().isoformat()

    return {
        "titolo": "Calendario (Eventi e Export ICS)",
        "sottotitolo": f"Revisione funzionale – generato il {today}",
        "blocchi": [
            (
                "Scopo",
                [
                    "Il calendario permette di registrare eventi (riunioni, assemblee, scadenze) e di esportarli in formato .ics (compatibile con Google/Outlook/Thunderbird).",
                ],
            ),
            (
                "Tipi evento",
                [
                    "Riunione CD, Assemblea ordinaria, Elezioni, Autorizzazione ponte, Scadenza assicurazione, Altro.",
                    "Ogni evento ha: tipo, titolo, start_ts (ISO datetime), descrizione, luogo, reminder_days, origin (collegamento logico).",
                ],
            ),
            (
                "Creazione e modifica (CalendarWizard)",
                [
                    "Dialog per creare/modificare singolo evento: selezione tipo, titolo, data/ora, luogo, promemoria (giorni prima), note.",
                    "Persistenza: add_calendar_event / update_calendar_event su DB.",
                ],
            ),
            (
                "Export in ICS",
                [
                    "events_to_ics(events): genera un VCALENDAR con VEVENT per ogni record.",
                    "Campi principali: UID deterministico librosoci-event-<id>@ari, DTSTART, SUMMARY, CATEGORIES, LOCATION, DESCRIPTION.",
                    "Export dalla UI: ‘Esporta tutto (.ics)’ e ‘Esporta selezionato (.ics)’.",
                ],
            ),
            (
                "Integrazioni automatiche",
                [
                    "Assicurazioni: alla creazione polizza viene creato un evento scadenza; in update può aggiornare la data.",
                    "Ponti: autorizzazioni possono generare/aggiornare eventi calendario per scadenze (origin collegato).",
                ],
            ),
        ],
    }


def _section_11_content() -> dict:
    today = date.today().isoformat()

    return {
        "titolo": "Gestione Duplicati (Ricerca e Merge)",
        "sottotitolo": f"Revisione funzionale – generato il {today}",
        "blocchi": [
            (
                "Scopo",
                [
                    "Questa funzione aiuta a individuare soci duplicati e a fonderli (merge) in un record master, spostando i documenti e tracciando l’evento.",
                ],
            ),
            (
                "Rilevazione duplicati",
                [
                    "Per matricola: soci con stessa matricola (case-insensitive, trimmed) e non in cestino.",
                    "Per nominativo: soci con stesso nominativo (case-insensitive, trimmed) e non in cestino.",
                    "Caso RCL: duplicati su coppia nome+cognome quando matricola/nominativo sono vuoti.",
                    "Combined: unisce i gruppi in componenti connesse (un record può collegare due gruppi diversi).",
                ],
            ),
            (
                "Merge (merge_duplicates)",
                [
                    "L’utente seleziona un MASTER e uno o più duplicati.",
                    "Aggiornamento master: applica solo i campi scelti (field_values).",
                    "Soft-delete duplicato: imposta soci.deleted_at = datetime('now').",
                    "Documenti: sposta i record in tabella documenti dal duplicato al master (UPDATE documenti SET socio_id=?).",
                    "Audit: inserisce un record in eventi_libro_soci con dettagli_json {type:'merge', merged_id, fields_updated}.",
                ],
            ),
            (
                "UI",
                [
                    "Dialog dedicato: selezione gruppo, confronto differenze campo-per-campo e conferma merge con report.",
                    "Dopo merge: refresh lista soci per riflettere l’unione.",
                ],
            ),
        ],
    }


def _section_12_content() -> dict:
    today = date.today().isoformat()

    return {
        "titolo": "Magazzino / Prestiti",
        "sottotitolo": f"Revisione funzionale – generato il {today}",
        "blocchi": [
            (
                "Scopo",
                [
                    "Gestione dell’inventario di sezione e dei prestiti agli associati: anagrafica oggetti, stato disponibilità, prestito e reso con storico.",
                ],
            ),
            (
                "Dati gestiti",
                [
                    "Oggetti: numero_inventario, marca, modello, descrizione, note.",
                    "Prestiti: socio_id, data_prestito, data_reso, note (storico per item).",
                    "Stato calcolato: ‘Disponibile’ vs ‘In prestito’ in base a prestito attivo (data_reso NULL).",
                ],
            ),
            (
                "Operazioni principali (logica)",
                [
                    "CRUD oggetti: create_item, update_item, delete_item.",
                    "Prestiti: create_loan con vincolo ‘un solo prestito attivo’ per oggetto; register_return per segnare il reso.",
                    "List: list_items include join sul prestito attivo e dati socio associato; list_loans mostra storico prestiti.",
                ],
            ),
            (
                "UI: MagazzinoPanel",
                [
                    "Toolbar: Nuovo oggetto, Salva, Elimina, Aggiorna; filtro per stato; ricerca full-text.",
                    "Lista oggetti: evidenzia oggetti in prestito.",
                    "Sezione prestiti: Nuovo prestito e Registra reso; mostra storico con socio/matricola e date.",
                ],
            ),
        ],
    }


def _section_13_content() -> dict:
    today = date.today().isoformat()

    return {
        "titolo": "Ponti / Ripetitori",
        "sottotitolo": f"Revisione funzionale – generato il {today}",
        "blocchi": [
            (
                "Scopo",
                [
                    "Gestione anagrafica ponti/ripetitori con stato, autorizzazioni/scadenze, documenti tecnici e promemoria calendario.",
                ],
            ),
            (
                "Dati e storage documenti",
                [
                    "Record ponte: nome, nominativo, località (qth), note tecniche, stato_corrente.",
                    "Autorizzazioni: tipo/ente/numero, date rilascio/scadenza, note, documento; collegate a un evento calendario.",
                    "Documenti ponte: file copiati in data/documents/ponti/<id>_<token>/<TIPO>/ con filename random (token hex).",
                ],
            ),
            (
                "Operazioni principali",
                [
                    "CRUD ponte: create_ponte, update_ponte, delete_ponte, list_ponti.",
                    "Autorizzazioni: save_authorization, list_authorizations, delete_authorization (con gestione eventi calendario).",
                    "Documenti: add_ponte_document, list_ponte_documents, update_ponte_document, delete_ponte_document.",
                ],
            ),
            (
                "UI: PontiPanel",
                [
                    "Elenco con filtro stato (ATTIVO/MANUTENZIONE/DISMESSO) e ricerca; evidenzia scadenze imminenti/scadute.",
                    "Blocchi: dettaglio ponte, autorizzazioni (con apri documento), documenti allegati (apri file/cartella).",
                ],
            ),
        ],
    }


def _section_14_content() -> dict:
    today = date.today().isoformat()

    return {
        "titolo": "Assicurazioni",
        "sottotitolo": f"Revisione funzionale – generato il {today}",
        "blocchi": [
            (
                "Scopo",
                [
                    "Gestione polizze assicurative della sezione: anagrafica polizza, scadenze con promemoria, pagamenti, coperture, documenti allegati.",
                ],
            ),
            (
                "Dati e automatismi",
                [
                    "Polizza: nome, compagnia, tipo_polizza, numero_polizza, data_inizio, data_scadenza, importo_annuale, stato, responsabile, note.",
                    "Alla creazione: genera un evento calendario ‘Scadenza assicurazione’ con reminder_days (default 60) e origin=assicurazione_<id>.",
                    "In update: se cambia data_scadenza aggiorna l’evento calendario collegato.",
                ],
            ),
            (
                "Pagamenti, coperture, documenti",
                [
                    "Pagamenti: add_pagamento/list_pagamenti/delete_pagamento (data, importo, metodo, riferimento, note).",
                    "Coperture: add_copertura/update_copertura/delete_copertura/list_coperture (massimali/franchigie/descrizione).",
                    "Documenti: add_assicurazione_document/list_assicurazione_documents/delete_assicurazione_document; storage in data/documents/assicurazioni/<id>_<nome>/.",
                ],
            ),
            (
                "UI: AssicurazioniPanel",
                [
                    "Elenco con filtro stato (ATTIVA/SCADUTA/DISDETTA/SOSPESA) e ricerca; evidenzia scadenze imminenti.",
                    "Blocchi: dettaglio polizza, pagamenti, coperture, documenti allegati (apri file/cartella).",
                ],
            ),
        ],
    }


SECTIONS: List[Section] = [
    Section(code="01", title="Panoramica e Avvio", file_stem="01_Panoramica_e_Avvio"),
    Section(code="02", title="Anagrafica Soci", file_stem="02_Anagrafica_Soci"),
    Section(code="03", title="Documenti Socio", file_stem="03_Documenti_Socio"),
    Section(code="04", title="Documenti di Sezione", file_stem="04_Documenti_Sezione"),
    Section(code="05", title="Backup e Ripristino", file_stem="05_Backup_e_Ripristino"),
    Section(code="06", title="Email Wizard", file_stem="06_Email_Wizard"),
    Section(code="07", title="Preferenze e Configurazione", file_stem="07_Preferenze_e_Config"),
    Section(code="08", title="Import/Export", file_stem="08_Import_Export"),
    Section(code="09", title="Consiglio Direttivo", file_stem="09_Consiglio_Direttivo"),
    Section(code="10", title="Calendario", file_stem="10_Calendario"),
    Section(code="11", title="Duplicati", file_stem="11_Duplicati"),
    Section(code="12", title="Magazzino", file_stem="12_Magazzino"),
    Section(code="13", title="Ponti", file_stem="13_Ponti"),
    Section(code="14", title="Assicurazioni", file_stem="14_Assicurazioni"),
]


CONTENT_BUILDERS = {
    "01": _section_01_content,
    "02": _section_02_content,
    "03": _section_03_content,
    "04": _section_04_content,
    "05": _section_05_content,
    "06": _section_06_content,
    "07": _section_07_content,
    "08": _section_08_content,
    "09": _section_09_content,
    "10": _section_10_content,
    "11": _section_11_content,
    "12": _section_12_content,
    "13": _section_13_content,
    "14": _section_14_content,
}


def _get_section(code: str) -> Section:
    for sec in SECTIONS:
        if sec.code == code:
            return sec
    raise SystemExit(f"Sezione non riconosciuta: {code}. Disponibili: {', '.join(s.code for s in SECTIONS)}")


def generate_section_docx(section_code: str) -> Path:
    try:
        from docx import Document  # type: ignore
    except Exception as exc:
        raise SystemExit(
            "python-docx non disponibile. Installa con: pip install python-docx\n"
            f"Dettagli: {exc}"
        )

    section = _get_section(section_code)
    OUTPUT_DIR.mkdir(parents=True, exist_ok=True)

    builder = CONTENT_BUILDERS.get(section_code)
    if not builder:
        raise SystemExit(
            f"Contenuto non implementato per sezione {section_code}. "
            f"Disponibili: {', '.join(sorted(CONTENT_BUILDERS.keys()))}"
        )
    content = builder()

    doc = Document()

    _add_heading(doc, f"GestioneSoci / GLR – {content['titolo']}", level=0)
    _add_paragraph(doc, content["sottotitolo"])

    for heading, paragraphs_or_bullets in content["blocchi"]:
        _add_heading(doc, heading, level=1)
        # Heuristic: if it's a list of short items => bullets; else paragraphs.
        if isinstance(paragraphs_or_bullets, list) and all(isinstance(x, str) for x in paragraphs_or_bullets):
            # Special-case blocks that intentionally include newlines.
            if any("\n" in x for x in paragraphs_or_bullets):
                for x in paragraphs_or_bullets:
                    _add_paragraph(doc, x)
            else:
                _add_bullets(doc, paragraphs_or_bullets)
        else:
            _add_paragraph(doc, str(paragraphs_or_bullets))

    out_path = OUTPUT_DIR / f"{section.file_stem}.docx"
    return _safe_doc_save(doc, out_path)


def generate_combined_docx() -> Path:
    """Generate a single .docx containing all sections in SECTIONS order."""

    try:
        from docx import Document  # type: ignore
    except Exception as exc:
        raise SystemExit(
            "python-docx non disponibile. Installa con: pip install python-docx\n"
            f"Dettagli: {exc}"
        )

    OUTPUT_DIR.mkdir(parents=True, exist_ok=True)
    today = date.today().isoformat()

    app_version, build_id = _read_version_build()

    doc = Document()

    # Cover page
    _add_centered_paragraph(doc, "GestioneSoci / GLR", style="Title")
    _add_centered_paragraph(doc, "Documentazione di revisione (completa)", style="Subtitle")
    _add_centered_paragraph(doc, " ")
    if app_version and build_id:
        _add_centered_paragraph(doc, f"Versione {app_version} – Build {build_id}")
    elif app_version:
        _add_centered_paragraph(doc, f"Versione {app_version}")
    elif build_id:
        _add_centered_paragraph(doc, f"Build {build_id}")
    _add_centered_paragraph(doc, f"Generato il {today}")
    doc.add_page_break()

    _add_heading(doc, "GestioneSoci / GLR – Documentazione di Revisione (completa)", level=0)
    _add_paragraph(doc, f"Generato il {today}")

    first = True
    for sec in SECTIONS:
        builder = CONTENT_BUILDERS.get(sec.code)
        if not builder:
            raise SystemExit(
                f"Contenuto non implementato per sezione {sec.code}. "
                f"Disponibili: {', '.join(sorted(CONTENT_BUILDERS.keys()))}"
            )
        content = builder()

        if not first:
            doc.add_page_break()
        first = False

        _add_heading(doc, f"Sezione {sec.code} – {content['titolo']}", level=0)
        _add_paragraph(doc, content["sottotitolo"])

        for heading, paragraphs_or_bullets in content["blocchi"]:
            _add_heading(doc, heading, level=1)
            if isinstance(paragraphs_or_bullets, list) and all(isinstance(x, str) for x in paragraphs_or_bullets):
                if any("\n" in x for x in paragraphs_or_bullets):
                    for x in paragraphs_or_bullets:
                        _add_paragraph(doc, x)
                else:
                    _add_bullets(doc, paragraphs_or_bullets)
            else:
                _add_paragraph(doc, str(paragraphs_or_bullets))

    out_path = OUTPUT_DIR / "00_Documentazione_Completa.docx"
    return _safe_doc_save(doc, out_path)


def main() -> int:
    parser = argparse.ArgumentParser(description="Generate review documentation (.docx) per sezione")
    group = parser.add_mutually_exclusive_group(required=True)
    group.add_argument(
        "--section",
        help="Codice sezione (es. 01)",
    )
    group.add_argument(
        "--all",
        action="store_true",
        help="Genera tutte le sezioni disponibili",
    )
    group.add_argument(
        "--combined",
        action="store_true",
        help="Genera un unico documento con tutte le sezioni",
    )
    args = parser.parse_args()

    if args.all:
        out_paths: list[Path] = []
        for sec in SECTIONS:
            out_paths.append(generate_section_docx(sec.code))

        # Also generate the combined document for convenience.
        out_paths.append(generate_combined_docx())

        for path in out_paths:
            print(str(path))
        print(f"Generated {len(out_paths)} documents in {OUTPUT_DIR}")
        return 0

    if args.combined:
        out_path = generate_combined_docx()
        print(str(out_path))
        return 0

    out_path = generate_section_docx(str(args.section))
    print(str(out_path))
    return 0


if __name__ == "__main__":
    raise SystemExit(main())
