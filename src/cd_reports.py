# -*- coding: utf-8 -*-
"""CD reporting / exports.

Currently supports exporting the "Libro verbali" as an Excel workbook.
"""

from __future__ import annotations

from dataclasses import dataclass
from typing import Iterable
from typing import Any, cast
import sqlite3
from pathlib import Path
from datetime import date


@dataclass(frozen=True)
class LibroVerbaliRow:
    numero: int
    data_iso: str
    odg: str


@dataclass(frozen=True)
class LibroDelibereRow:
    numero_riga: int
    data_iso: str
    numero_delibera: str
    oggetto: str
    esito: str
    note: str
    favorevoli: int | None
    contrari: int | None
    astenuti: int | None


def _odg_json_to_text(odg_json: str | None) -> str:
    if not odg_json:
        return ""
    try:
        import json

        payload = json.loads(odg_json)
        items = payload.get("items")
        if not isinstance(items, list):
            return ""
        titles: list[str] = []
        for item in items:
            if isinstance(item, dict):
                title = str(item.get("title") or "").strip()
                if not title:
                    continue
                if bool(item.get("requires_delibera")):
                    titles.append(f"[D] {title}")
                else:
                    titles.append(title)
        return "\n".join(titles)
    except Exception:
        return ""


def _iter_libro_verbali_rows() -> list[LibroVerbaliRow]:
    """Load meetings from DB and normalize them into rows for the book."""

    from database import fetch_all
    from utils import today_iso

    cutoff = today_iso()

    rows = fetch_all(
        """
        SELECT id, data, note, odg_json, tipo_riunione
        FROM cd_riunioni
        WHERE data IS NOT NULL
          AND TRIM(data) <> ''
          AND LOWER(COALESCE(tipo_riunione, 'passata')) <> 'futura'
          AND data <= ?
        ORDER BY data ASC, id ASC
        """,
        (cutoff,),
    )

    result: list[LibroVerbaliRow] = []
    counter = 0
    for row in rows:
        meeting = dict(row)
        data_iso = str(meeting.get("data") or "").strip()
        if not data_iso:
            continue

        odg_text = (str(meeting.get("note") or "").strip() if meeting.get("note") is not None else "")
        if not odg_text:
            odg_text = _odg_json_to_text(meeting.get("odg_json"))

        counter += 1
        result.append(LibroVerbaliRow(numero=counter, data_iso=data_iso, odg=odg_text or ""))

    return result


def _get_cd_delibere_date_expr(conn: sqlite3.Connection) -> str:
    """Return a SQL expression for the delibera date column.

    Older databases may miss `data_votazione` or use a legacy `data` column.
    Returned expression is safe to embed as-is (it includes the `d.` prefix).
    """

    try:
        cols = {str(r["name"]).lower() for r in conn.execute("PRAGMA table_info(cd_delibere)")}
    except Exception:
        cols = set()
    if "data_votazione" in cols:
        return "d.data_votazione"
    if "data" in cols:
        return "d.data"
    return "NULL"


def _iter_libro_delibere_rows() -> list[LibroDelibereRow]:
    """Load delibere from DB and normalize them into rows for the book."""

    from database import fetch_all, get_connection
    from utils import today_iso

    cutoff = today_iso()

    try:
        with get_connection() as conn:
            date_expr = _get_cd_delibere_date_expr(conn)
    except Exception:
        date_expr = "d.data_votazione"

    rows = fetch_all(
        f"""
        SELECT
            d.id,
            COALESCE({date_expr}, r.data) AS data_iso,
            d.numero AS numero_delibera,
            d.oggetto,
            d.esito,
                        d.note,
                        d.favorevoli,
                        d.contrari,
                        d.astenuti,
            r.data AS data_riunione
        FROM cd_delibere d
        JOIN cd_riunioni r ON r.id = d.cd_id
        WHERE r.data IS NOT NULL
          AND TRIM(r.data) <> ''
          AND LOWER(COALESCE(r.tipo_riunione, 'passata')) <> 'futura'
          AND r.data <= ?
        ORDER BY data_iso ASC, d.id ASC
        """,
        (cutoff,),
    )

    result: list[LibroDelibereRow] = []
    counter = 0
    for row in rows:
        d = dict(row)
        data_iso = str(d.get("data_iso") or "").strip() or str(d.get("data_riunione") or "").strip()
        if not data_iso:
            continue
        counter += 1
        result.append(
            LibroDelibereRow(
                numero_riga=counter,
                data_iso=data_iso,
                numero_delibera=str(d.get("numero_delibera") or "").strip(),
                oggetto=str(d.get("oggetto") or "").strip(),
                esito=str(d.get("esito") or "").strip(),
                note=str(d.get("note") or "").strip(),
                favorevoli=(int(d["favorevoli"]) if d.get("favorevoli") is not None else None),
                contrari=(int(d["contrari"]) if d.get("contrari") is not None else None),
                astenuti=(int(d["astenuti"]) if d.get("astenuti") is not None else None),
            )
        )

    return result


def export_libro_verbali_xlsx(output_path: str) -> tuple[int, list[str]]:
    """Export "Libro verbali" to an Excel (.xlsx) file.

    Layout matches the provided template screenshot:
    - Columns: N. | data | odg
    - ODG wrapped with borders.

    Returns:
        (written_rows, warnings)
    """

    warnings: list[str] = []

    try:
        import openpyxl
        from openpyxl.styles import Alignment, Border, Font, Side
    except Exception as exc:  # pragma: no cover
        return 0, [f"openpyxl non disponibile: {exc}"]

    from utils import iso_to_ddmmyyyy

    data_rows = _iter_libro_verbali_rows()

    wb = openpyxl.Workbook()
    ws = cast(Any, wb.active)
    ws.title = "Foglio1"

    # Header
    ws["A1"].value = "N."
    ws["B1"].value = "data"
    ws["C1"].value = "odg"

    header_font = Font(bold=True)
    for cell in (ws["A1"], ws["B1"], ws["C1"]):
        cell.font = header_font

    # Column widths (approximate the screenshot)
    ws.column_dimensions["A"].width = 4.5
    ws.column_dimensions["B"].width = 12
    ws.column_dimensions["C"].width = 110

    wrap_top = Alignment(wrap_text=True, vertical="top")
    center = Alignment(vertical="top", horizontal="center", wrap_text=True)

    thin = Side(style="thin")
    border = Border(left=thin, right=thin, top=thin, bottom=thin)

    # Freeze header
    ws.freeze_panes = "A2"

    row_idx = 2
    for item in data_rows:
        ws.cell(row=row_idx, column=1, value=item.numero)
        ws.cell(row=row_idx, column=2, value=iso_to_ddmmyyyy(item.data_iso))
        ws.cell(row=row_idx, column=3, value=item.odg)

        ws.cell(row=row_idx, column=1).alignment = center
        ws.cell(row=row_idx, column=2).alignment = center
        ws.cell(row=row_idx, column=3).alignment = wrap_top

        # Apply border
        for col in (1, 2, 3):
            c = ws.cell(row=row_idx, column=col)
            c.border = border

        row_idx += 1

    # Apply border to header too
    for col in (1, 2, 3):
        c = ws.cell(row=1, column=col)
        c.border = border

    try:
        wb.save(output_path)
    except Exception as exc:
        return 0, [f"Impossibile salvare il file: {exc}"]

    return len(data_rows), warnings


def export_libro_delibere_xlsx(output_path: str) -> tuple[int, list[str]]:
    """Export "Libro delibere" to an Excel (.xlsx) file.

    Layout (minimal, consistent with existing UI fields):
    - Columns: N. | data | numero | oggetto | esito

    Returns:
        (written_rows, warnings)
    """

    warnings: list[str] = []

    try:
        import openpyxl
        from openpyxl.styles import Alignment, Border, Font, Side
    except Exception as exc:  # pragma: no cover
        return 0, [f"openpyxl non disponibile: {exc}"]

    from utils import iso_to_ddmmyyyy

    data_rows = _iter_libro_delibere_rows()

    wb = openpyxl.Workbook()
    ws = cast(Any, wb.active)
    ws.title = "Foglio1"

    ws["A1"].value = "N."
    ws["B1"].value = "data"
    ws["C1"].value = "numero"
    ws["D1"].value = "oggetto"
    ws["E1"].value = "esito"

    header_font = Font(bold=True)
    for cell in (ws["A1"], ws["B1"], ws["C1"], ws["D1"], ws["E1"]):
        cell.font = header_font

    ws.column_dimensions["A"].width = 4.5
    ws.column_dimensions["B"].width = 12
    ws.column_dimensions["C"].width = 14
    ws.column_dimensions["D"].width = 90
    ws.column_dimensions["E"].width = 14

    wrap_top = Alignment(wrap_text=True, vertical="top")
    center = Alignment(vertical="top", horizontal="center", wrap_text=True)

    thin = Side(style="thin")
    border = Border(left=thin, right=thin, top=thin, bottom=thin)

    ws.freeze_panes = "A2"

    row_idx = 2
    for item in data_rows:
        ws.cell(row=row_idx, column=1, value=item.numero_riga)
        ws.cell(row=row_idx, column=2, value=iso_to_ddmmyyyy(item.data_iso))
        ws.cell(row=row_idx, column=3, value=item.numero_delibera)
        ws.cell(row=row_idx, column=4, value=item.oggetto)
        ws.cell(row=row_idx, column=5, value=item.esito)

        ws.cell(row=row_idx, column=1).alignment = center
        ws.cell(row=row_idx, column=2).alignment = center
        ws.cell(row=row_idx, column=3).alignment = center
        ws.cell(row=row_idx, column=4).alignment = wrap_top
        ws.cell(row=row_idx, column=5).alignment = center

        for col in (1, 2, 3, 4, 5):
            c = ws.cell(row=row_idx, column=col)
            c.border = border

        row_idx += 1

    for col in (1, 2, 3, 4, 5):
        c = ws.cell(row=1, column=col)
        c.border = border

    try:
        wb.save(output_path)
    except Exception as exc:
        return 0, [f"Impossibile salvare il file: {exc}"]

    return len(data_rows), warnings


def _clear_docx_body(doc: Any) -> None:
    """Remove all body content from a python-docx Document, keeping section properties."""

    body = doc._element.body  # type: ignore[attr-defined]
    for child in list(body):
        if child.tag.endswith("}sectPr"):
            continue
        body.remove(child)


def export_libro_delibere_docx(
    output_path: str,
    *,
    template_path: str | None = None,
) -> tuple[int, list[str]]:
    """Export "Libro delibere" to a Word (.docx) file.

    This exporter is template-driven and paragraph-based (the provided ARI-BG template
    contains no tables). It keeps template styles and regenerates the content:
    - Title line: "Delibere"
    - Second line: year range (e.g. "2023-2025") derived from exported rows
    - Then for each delibera:
        "delibera n. XX/YYYY (Oggetto)" + detail paragraphs.

    Returns:
        (written_delibere, warnings)
    """

    warnings: list[str] = []

    try:
        from docx import Document
    except Exception as exc:  # pragma: no cover
        return 0, [f"python-docx non disponibile: {exc}"]

    rows = _iter_libro_delibere_rows()

    doc: Any
    if template_path and Path(template_path).exists():
        doc = Document(template_path)
    else:
        doc = Document()
        if template_path:
            warnings.append(f"Template non trovato: {template_path}")

    _clear_docx_body(doc)

    # Year range from exported rows
    years: list[int] = []
    for r in rows:
        try:
            years.append(int(str(r.data_iso).split("-")[0]))
        except Exception:
            continue
    if years:
        y1, y2 = min(years), max(years)
        year_range = f"{y1}-{y2}" if y1 != y2 else f"{y1}"
    else:
        year_range = str(date.today().year)

    doc.add_paragraph("Delibere")
    doc.add_paragraph(year_range)
    doc.add_paragraph("")

    for r in rows:
        title = f"delibera n. {r.numero_delibera} ({r.oggetto})".strip()
        doc.add_paragraph(title)

        if r.note:
            doc.add_paragraph(r.note)

        # Compact outcome line (if available)
        parts: list[str] = []
        if r.esito:
            parts.append(r.esito)
        if r.favorevoli is not None or r.contrari is not None or r.astenuti is not None:
            fv = "" if r.favorevoli is None else str(r.favorevoli)
            cn = "" if r.contrari is None else str(r.contrari)
            asn = "" if r.astenuti is None else str(r.astenuti)
            parts.append(f"Favorevoli: {fv}  Contrari: {cn}  Astenuti: {asn}".strip())
        if parts:
            doc.add_paragraph(" - ".join([p for p in parts if p]))

    try:
        doc.save(output_path)
    except Exception as exc:
        return 0, [f"Impossibile salvare il file: {exc}"]

    return len(rows), warnings


def _normalize_cd_carica_codice(nome: str) -> str:
    try:
        from database import _normalize_cd_carica_codice as _n

        return _n(nome)
    except Exception:
        return str(nome or "").strip().lower()


def _pick_first_by_role(composizione: list[dict[str, Any]], role_codice: str) -> str:
    role_codice_n = _normalize_cd_carica_codice(role_codice)
    for item in composizione or []:
        try:
            carica = str(item.get("carica") or "")
            nome = str(item.get("nome") or "").strip()
        except Exception:
            continue
        if not nome:
            continue
        if _normalize_cd_carica_codice(carica) == role_codice_n:
            return nome
    return ""


def _odg_for_meeting(note: str | None, odg_json: str | None) -> str:
    text = (str(note or "").strip() if note is not None else "")
    if text:
        return text
    return _odg_json_to_text(odg_json)


def _replace_in_docx_text(doc: Any, replacements: dict[str, str]) -> None:
    """Replace placeholders in a python-docx Document.

    Supports both `{KEY}` and `[KEY]` placeholders.
    Note: replacement is run-based; placeholders should be contained within a single run
    in the template for reliable substitution.
    """

    def _replace_in_paragraph(p: Any) -> None:
        if not getattr(p, "runs", None):
            return
        for run in p.runs:
            t = run.text
            if not t:
                continue
            for k, v in replacements.items():
                if not k:
                    continue
                t = t.replace(f"{{{k}}}", v)
                t = t.replace(f"[{k}]", v)
            run.text = t

    def _replace_in_table(table: Any) -> None:
        for row in getattr(table, "rows", []) or []:
            for cell in getattr(row, "cells", []) or []:
                for p in getattr(cell, "paragraphs", []) or []:
                    _replace_in_paragraph(p)
                for t2 in getattr(cell, "tables", []) or []:
                    _replace_in_table(t2)

    for p in getattr(doc, "paragraphs", []) or []:
        _replace_in_paragraph(p)
    for table in getattr(doc, "tables", []) or []:
        _replace_in_table(table)

    # Headers/footers
    for section in getattr(doc, "sections", []) or []:
        try:
            header = section.header
            for p in getattr(header, "paragraphs", []) or []:
                _replace_in_paragraph(p)
            for table in getattr(header, "tables", []) or []:
                _replace_in_table(table)
        except Exception:
            pass
        try:
            footer = section.footer
            for p in getattr(footer, "paragraphs", []) or []:
                _replace_in_paragraph(p)
            for table in getattr(footer, "tables", []) or []:
                _replace_in_table(table)
        except Exception:
            pass


def export_verbale_cd_docx(
    meeting_id: int,
    output_path: str,
    *,
    template_path: str | None = None,
) -> tuple[bool, list[str]]:
    """Generate a single CD verbale as a Word (.docx) file.

    The output is template-driven if `template_path` is provided.
    Placeholders supported in the template (either `{KEY}` or `[KEY]`):
    - DATA_ISO, DATA
    - NUMERO_CD, TITOLO
    - ODG
    - PRESIDENTE, SEGRETARIO
    - CD_COMPOSIZIONE
    - DELIBERE
    """

    warnings: list[str] = []

    try:
        from docx import Document
    except Exception as exc:  # pragma: no cover
        return False, [f"python-docx non disponibile: {exc}"]

    from database import fetch_one, fetch_all
    from utils import iso_to_ddmmyyyy

    meeting = fetch_one(
        """
        SELECT id, data, numero_cd, titolo, note, odg_json
        FROM cd_riunioni
        WHERE id = ?
        """,
        (int(meeting_id),),
    )
    if not meeting:
        return False, ["Riunione non trovata"]

    m = dict(meeting)
    data_iso = str(m.get("data") or "").strip()
    numero_cd = str(m.get("numero_cd") or "").strip()
    titolo = str(m.get("titolo") or "").strip()
    odg = _odg_for_meeting(m.get("note"), m.get("odg_json"))

    # CD composition for the meeting (mandate-aware)
    try:
        from cd_mandati import get_cd_composizione_for_meeting

        composizione = get_cd_composizione_for_meeting(int(meeting_id))
    except Exception as exc:
        composizione = []
        warnings.append(f"Impossibile caricare composizione CD: {exc}")

    presidente = _pick_first_by_role(composizione, "presidente")
    segretario = _pick_first_by_role(composizione, "segretario")

    comp_lines: list[str] = []
    for item in composizione or []:
        carica = str(item.get("carica") or "").strip()
        nome = str(item.get("nome") or "").strip()
        if not (carica and nome):
            continue
        comp_lines.append(f"{carica}: {nome}")
    cd_composizione_text = "\n".join(comp_lines)

    # Delibere for the meeting (best-effort)
    delibere_lines: list[str] = []
    try:
        delibere = fetch_all(
            """
            SELECT numero, oggetto, esito, note, favorevoli, contrari, astenuti
            FROM cd_delibere
            WHERE cd_id = ?
            ORDER BY id ASC
            """,
            (int(meeting_id),),
        )
        for drow in delibere or []:
            d = dict(drow)
            num = str(d.get("numero") or "").strip()
            ogg = str(d.get("oggetto") or "").strip()
            esito = str(d.get("esito") or "").strip()
            line = " - ".join([p for p in (num, ogg, esito) if p])
            if line:
                delibere_lines.append(line)
    except Exception as exc:
        warnings.append(f"Impossibile caricare delibere: {exc}")
    delibere_text = "\n".join(delibere_lines)

    repl: dict[str, str] = {
        "DATA_ISO": data_iso,
        "DATA": iso_to_ddmmyyyy(data_iso) if data_iso else "",
        "NUMERO_CD": numero_cd,
        "TITOLO": titolo,
        "ODG": odg,
        "PRESIDENTE": presidente,
        "SEGRETARIO": segretario,
        "CD_COMPOSIZIONE": cd_composizione_text,
        "DELIBERE": delibere_text,
    }

    doc: Any
    if template_path and Path(template_path).exists():
        doc = Document(template_path)
        _replace_in_docx_text(doc, repl)
    else:
        if template_path:
            warnings.append(f"Template non trovato: {template_path}")
        doc = Document()
        doc.add_paragraph("Verbale Consiglio Direttivo")
        header_bits: list[str] = []
        if numero_cd:
            header_bits.append(f"N. {numero_cd}")
        if data_iso:
            header_bits.append(repl["DATA"])
        if header_bits:
            doc.add_paragraph(" - ".join(header_bits))
        if titolo:
            doc.add_paragraph(titolo)
        doc.add_paragraph("")
        if presidente or segretario:
            if presidente:
                doc.add_paragraph(f"Presidente: {presidente}")
            if segretario:
                doc.add_paragraph(f"Segretario: {segretario}")
            doc.add_paragraph("")
        if cd_composizione_text:
            doc.add_paragraph("Composizione CD:")
            for line in comp_lines:
                doc.add_paragraph(line)
            doc.add_paragraph("")
        if odg:
            doc.add_paragraph("Ordine del Giorno:")
            for line in odg.splitlines():
                doc.add_paragraph(line)
            doc.add_paragraph("")
        if delibere_text:
            doc.add_paragraph("Delibere:")
            for line in delibere_lines:
                doc.add_paragraph(line)

    try:
        doc.save(output_path)
    except Exception as exc:
        return False, [f"Impossibile salvare il file: {exc}"]

    return True, warnings
