# -*- coding: utf-8 -*-
"""
Generazione libro delle delibere esportabile in PDF/DOCX, raggruppato per triennio di mandato.
"""
import os
from datetime import datetime
from cd_delibere import get_all_delibere
from cd_mandati import get_all_cd_mandati
from docx import Document
from docx.shared import Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH

# Helper: determina triennio da data
def get_triennio_for_date(date_str):
    try:
        year = int(date_str[:4])
        start = year - (year % 3)
        end = start + 2
        return f"{start}-{end}"
    except Exception:
        return "?"

# Raggruppa delibere per triennio
def group_delibere_by_triennio(delibere):
    grouped = {}
    for d in delibere:
        data = d.get('data_votazione') or d.get('created_at')
        triennio = get_triennio_for_date(data)
        grouped.setdefault(triennio, []).append(d)
    return grouped

# Genera DOCX libro delle delibere
def export_libro_delibere_docx(output_path):
    delibere = get_all_delibere()
    grouped = group_delibere_by_triennio(delibere)
    doc = Document()
    doc.add_heading('Libro delle Delibere', 0)
    for triennio, items in sorted(grouped.items()):
        doc.add_heading(f"Mandato {triennio}", level=1)
        table = doc.add_table(rows=1, cols=8)
        hdr_cells = table.rows[0].cells
        hdr_cells[0].text = 'Numero'
        hdr_cells[1].text = 'Data'
        hdr_cells[2].text = 'Oggetto'
        hdr_cells[3].text = 'Esito'
        hdr_cells[4].text = 'Fav.'
        hdr_cells[5].text = 'Contr.'
        hdr_cells[6].text = 'Asten.'
        hdr_cells[7].text = 'Verbale Rif.'
        for d in items:
            row = table.add_row().cells
            row[0].text = str(d.get('numero',''))
            row[1].text = str(d.get('data_votazione',''))
            row[2].text = str(d.get('oggetto',''))
            row[3].text = str(d.get('esito',''))
            row[4].text = str(d.get('favorevoli',''))
            row[5].text = str(d.get('contrari',''))
            row[6].text = str(d.get('astenuti',''))
            row[7].text = str(d.get('verbale_riferimento',''))
        doc.add_paragraph()
    doc.save(output_path)
    return output_path


# Esporta libro delle delibere in PDF
from fpdf import FPDF
def export_libro_delibere_pdf(output_path):
    delibere = get_all_delibere()
    grouped = group_delibere_by_triennio(delibere)
    pdf = FPDF()
    pdf.set_auto_page_break(auto=True, margin=15)
    pdf.add_page()
    pdf.set_font("Arial", 'B', 16)
    pdf.cell(0, 10, "Libro delle Delibere", ln=True, align='C')
    pdf.set_font("Arial", '', 12)
    for triennio, items in sorted(grouped.items()):
        pdf.ln(8)
        pdf.set_font("Arial", 'B', 14)
        pdf.cell(0, 8, f"Mandato {triennio}", ln=True)
        pdf.set_font("Arial", 'B', 10)
        headers = ['Numero', 'Data', 'Oggetto', 'Esito', 'Fav.', 'Contr.', 'Asten.', 'Verbale Rif.']
        col_widths = [20, 22, 60, 18, 12, 12, 12, 30]
        for i, h in enumerate(headers):
            pdf.cell(col_widths[i], 7, h, border=1, align='C')
        pdf.ln()
        pdf.set_font("Arial", '', 9)
        for d in items:
            row = [
                str(d.get('numero','')),
                str(d.get('data_votazione','')),
                str(d.get('oggetto','')),
                str(d.get('esito','')),
                str(d.get('favorevoli','')),
                str(d.get('contrari','')),
                str(d.get('astenuti','')),
                str(d.get('verbale_riferimento','')),
            ]
            for i, val in enumerate(row):
                pdf.cell(col_widths[i], 6, val[:40], border=1)
            pdf.ln()
    pdf.output(output_path)
    return output_path

if __name__ == "__main__":
    out_docx = export_libro_delibere_docx("libro_delibere_test.docx")
    print(f"Libro delle delibere esportato DOCX: {out_docx}")
    out_pdf = export_libro_delibere_pdf("libro_delibere_test.pdf")
    print(f"Libro delle delibere esportato PDF: {out_pdf}")
